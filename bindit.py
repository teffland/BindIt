import docx
import re

# Generate Regular Experessions
class Searcher():
    def __init__(self):
        self.settings = self.read_settings('settings.txt')
        self.input_filenames = self.get_input_filenames(self.settings['input_file'])
        self.sequences = self.get_sequence_expressions(self.settings['sequence_file'])
        self.regexs = []
        for sequence in self.sequences:
            self.regexs.extend(self.gen_regex(sequence, permute=self.settings['permute']))

    def read_settings(self, settings_file):
        lines = open(settings_file).read().splitlines()
        lines = [ line.strip(' ').split(':')[1].strip() for line in lines
                  if line and not line.strip(' ')[0] == '#']
        settings = {'input_file': lines[0],
                    'sequence_file': lines[1],
                    'scan_prefix': lines[2],
                    'permute': self.string_to_bool(lines[3]),
                    'verbosity': self.string_to_int(lines[4]) 
                    }
        return settings
        
    def get_input_filenames(self, input_file):
        input_filenames = [ filename.strip() for filename in open(input_file).read().splitlines() 
                            if filename and not filename.strip()[0] == '#']
        return input_filenames

    def get_sequence_expressions(self, sequence_file):
        sequences = [ sequence.strip(' ').split(',') for sequence in open(sequence_file).read().splitlines() 
                            if sequence and not sequence.strip(' ')[0] == '#']
        return sequences

    def gen_regex(self, sequence, permute=False):
        """Take a simple list of sequences representing the laymen way to write a regular expression.
        
        Sequences are of the format: [nucleotide sequence, maximum gap length, nucleodtide sequence, max gap length, ...]
            
            eg, ['TCTG', '4', 'CAGA']
        
        This simple way of wtriting a potential sequence will be translated to the appropriate regular expression.
        
            eg, 'TCTG{1}[AGTC]{0,4}CAGA{1}'
            
        
        Permutations are whether we should generate a regular expression for each possible regular expression
        that is one edit-distance from the provided sequence.
        """
        nucs = set(['A','C','G','T'])
        wildcard = "[ACGT]"
        reglist = []
        if permute:
            pass
            # TODO: Design permutations algorithm with Brett's constraints
            """
            n = len([s for s in sequence if type(s)==str]) # number of letters in total sequence
            for i in range(n):
                reg = ''
                for s in sequence:
                    if type(s) == str:
                        reg += s+'{1}'
                    if type(s) == int:
                        reg += wildcard+'{0,'+str(s)+'}'
                # loop over the regex, checking for the right spot to insert a wildcard
                j = 0        
                for c in reg:
                    if c in nucs:
                        if i == j:
                            c = wildcard+'{1}'
                            break
                        j += 1
                reglist.append(reg)
            """         
        else:
            reg = ''
            for s in sequence:
                try: s = int(s)
                except:
                    pass
                if type(s) == str:
                    reg += s+'{1}'
                if type(s) == int:
                    reg += wildcard+'{0,'+str(s)+'}'
            reglist.append(reg)
        return reglist
            
    def analyze_file(self, file_name, regex_list, scan_prefix):
        # TODO: Make this process more sophisticated
        # scan the doc
        doc = docx.Document(file_name)
        data = ''.join([p.text for p in doc.paragraphs]) # convert doc into a big string
        for reg in regex_list: # scan the document, splicing out matching sequences
            split = zip(re.split(reg,data), re.findall(reg, data))

        # write out an edited version
        scanned = docx.Document('template.docx')
        p = scanned.add_paragraph('')
        for bad,good in split:
            # the inbetween sequences
            run = p.add_run(bad)
            run.font.name = "Courier New"
            run.font.size = docx.shared.Pt(10)
            # the tagged sequences
            run = p.add_run(good)
            run.bold = True
            run.font.name = "Courier New"
            run.font.size = docx.shared.Pt(14)
        scanned.save(scan_prefix+'_'+file_name)

    def analyze_all(self):
        for infile in self.input_filenames:
            print "Analyzing %s" % infile
            self.analyze_file(infile, self.regexs, self.settings['scan_prefix'])

    def string_to_bool(self, string):
        if string.lower() in ['true', 't', 'yes', 'fursure', 'yeah', 'yea']: return True
        return False

    def string_to_int(self, string):
        try: return int(s)
        except: return 0

# run the program
main = Searcher()
main.analyze_all()